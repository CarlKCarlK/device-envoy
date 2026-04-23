#!/usr/bin/env python3
"""Export a rustdoc HTML directory into a single DOCX file.

Usage:
  python3 scripts/rustdoc_site_to_docx.py \
    --site-dir target/riscv32imac-unknown-none-elf/doc/device_envoy_esp \
    --output docs/device_envoy_esp_site.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document


HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
BLOCK_CONTAINER_TAGS = {
    "article",
    "section",
    "div",
    "details",
    "summary",
    "main",
}
SKIP_TAGS = {"script", "style", "nav", "noscript"}


def discover_html_pages(site_dir: Path) -> list[Path]:
    pages = sorted(site_dir.rglob("*.html"))
    index_path = site_dir / "index.html"
    if index_path in pages:
        pages.remove(index_path)
        pages.insert(0, index_path)
    return pages


def extract_main(soup: BeautifulSoup) -> Tag | None:
    main = soup.find("main")
    if main:
        return main
    article = soup.find("article")
    if article:
        return article
    body = soup.find("body")
    if body:
        return body
    return None


def clean_text(value: str) -> str:
    return " ".join(value.split())


def tag_text(tag: Tag) -> str:
    return clean_text(tag.get_text(" ", strip=True))


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"


def render_list(document: Document, tag: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for list_item in tag.find_all("li", recursive=False):
        text = tag_text(list_item)
        if text:
            document.add_paragraph(text, style=style)
        # Render nested lists explicitly after the item line.
        for nested_list in list_item.find_all(["ul", "ol"], recursive=False):
            render_list(document, nested_list, ordered=(nested_list.name == "ol"))


def render_table(document: Document, tag: Tag) -> None:
    rows = tag.find_all("tr")
    if not rows:
        return

    column_count = max(len(row.find_all(["th", "td"])) for row in rows)
    if column_count == 0:
        return

    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for column_index, cell in enumerate(cells):
            table.cell(row_index, column_index).text = tag_text(cell)


def render_blocks(document: Document, element: Tag) -> None:
    for child in element.children:
        if isinstance(child, NavigableString):
            text = clean_text(str(child))
            if text:
                document.add_paragraph(text)
            continue

        if not isinstance(child, Tag):
            continue

        if child.name in SKIP_TAGS:
            continue

        if child.name in HEADING_TAGS:
            level = int(child.name[1])
            text = tag_text(child)
            if text:
                document.add_heading(text, level=min(level, 4))
            continue

        if child.name == "p":
            text = tag_text(child)
            if text:
                document.add_paragraph(text)
            continue

        if child.name == "pre":
            text = child.get_text("\n", strip=True)
            if text:
                add_code_block(document, text)
            continue

        if child.name == "blockquote":
            text = tag_text(child)
            if text:
                document.add_paragraph(text, style="Intense Quote")
            continue

        if child.name == "ul":
            render_list(document, child, ordered=False)
            continue

        if child.name == "ol":
            render_list(document, child, ordered=True)
            continue

        if child.name == "table":
            render_table(document, child)
            continue

        if child.name in BLOCK_CONTAINER_TAGS:
            render_blocks(document, child)
            continue

        # Fallback for uncommon tags that still contain readable text.
        text = tag_text(child)
        if text:
            document.add_paragraph(text)


def page_title(main: Tag, path: Path) -> str:
    heading = main.find("h1")
    if heading:
        title_text = tag_text(heading)
        if title_text:
            return title_text
    return path.name


def export_site(site_dir: Path, output_path: Path) -> tuple[int, int]:
    pages = discover_html_pages(site_dir)
    document = Document()

    document.add_heading(f"Rustdoc Export: {site_dir.name}", level=0)
    document.add_paragraph(f"Source directory: {site_dir}")
    document.add_paragraph(f"Page count: {len(pages)}")

    rendered = 0
    skipped = 0

    for page in pages:
        html = page.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "lxml")
        main = extract_main(soup)
        if main is None:
            skipped += 1
            continue

        rel_path = page.relative_to(site_dir)
        title = page_title(main, page)
        document.add_page_break()
        document.add_heading(f"{title}", level=1)
        document.add_paragraph(f"Path: {rel_path}")

        render_blocks(document, main)
        rendered += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return rendered, skipped


def parse_args(argv: Iterable[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--site-dir",
        required=True,
        type=Path,
        help="Path to rustdoc crate directory (contains index.html)",
    )
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Path to output DOCX file",
    )
    return parser.parse_args(argv)


def main() -> int:
    args = parse_args()
    site_dir = args.site_dir.resolve()
    output_path = args.output.resolve()

    index_path = site_dir / "index.html"
    if not index_path.exists():
        raise SystemExit(f"Expected rustdoc index not found: {index_path}")

    rendered, skipped = export_site(site_dir, output_path)
    print(f"Wrote DOCX: {output_path}")
    print(f"Rendered pages: {rendered}")
    print(f"Skipped pages: {skipped}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
